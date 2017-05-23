#! /usr/local/bin/python2
import PyPDF2
import docx
import os
import glob
import csv
import re
from indexer import *


#sample for extracting text from first page of file
#f=open("../data/gandhinagar.pdf","rb")
#print PyPDF2.PdfFileReader(f).getPage(0).extractText()


os.chdir("./data")

#pdffiles=[]
#for file in os.listdir('../data/'):
#	if file.endswith(".pdf"):
#		pdffiles+=[file]
#		f=open("../data/"+file,"rb")
#		pdf=PyPDF2.PdfFileReader(f)
#		n=pdf.numPages
#		for i in range(n):
#			print pdf.getPage(i).extractText().encode("utf8")

dict={"aplle":"a"}

for root,directories,filenames in os.walk("."):
	for directory in directories:
		print os.path.join(root,directory)
	for filename in filenames:
		cstr=''
		file= os.path.join(root,filename)
		if file.endswith(".pdf"):
			f=open(file,"rb")
			pdf=PyPDF2.PdfFileReader(f)
			n=pdf.numPages
			for i in range(n):
				t= pdf.getPage(i).extractText().encode("ascii","ignore").split()
				cstr+=' '.join(t)
		elif file.endswith(".docx"):
			doc = docx.Document(file)
			n=len(doc.paragraphs)
			for i in range(n):
				cstr+= doc.paragraphs[i].text
		elif file.endswith(".csv"):
			f=open(file)
			reader=csv.reader(f)
			for row in reader:
				cstr+=' '.join(row)
				
		elif file.endswith(".txt"):
			cstr+= open(file,"r").read()
		elif file.endswith(".c"):
			cstr+= open(file,"r").read()
		elif file.endswith(".cpp"):
			cstr+= open(file,"r").read()	
		elif file.endswith(".py"):
			cstr+= open(file,"r").read()
		if len(cstr)!=0:
			# Call the indexing function
			
			index(cstr)
			cstr=''



